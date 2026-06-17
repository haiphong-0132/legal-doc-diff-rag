import tempfile
import os
from pathlib import Path
from typing import Dict, Any
import datetime

import docx
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def format_id(chunk_id: str) -> str:
    if not chunk_id:
        return "Không rõ"
    try:
        from src.core.embedding import decode_section_id
        return decode_section_id(chunk_id)
    except:
        return chunk_id

def build_docx_report(data: Dict[str, Any]) -> Path:
    doc = docx.Document()
    
    title = doc.add_heading('BÁO CÁO SO SÁNH VĂN BẢN PHÁP LÝ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Ngày tạo: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('THỐNG KÊ TỔNG QUAN', level=1)
    stats = data.get("stats", {})
    stat_p = doc.add_paragraph()
    stat_p.add_run(f"Số lượng chunk VB1: {stats.get('so_luong_chunk_vb1', 0)}\n")
    stat_p.add_run(f"Số lượng chunk VB2: {stats.get('so_luong_chunk_vb2', 0)}\n")
    stat_p.add_run(f"Giống hệt: {stats.get('giong_nhau_hoan_toan', 0)}\n")
    stat_p.add_run(f"Sửa đổi: {stats.get('sua_doi', 0)}\n")
    stat_p.add_run(f"Thêm mới: {stats.get('them_moi', 0)}\n")
    stat_p.add_run(f"Xóa bỏ: {stats.get('xoa_bo', 0)}\n")
    if "elapsed_s" in stats and stats["elapsed_s"] is not None:
        stat_p.add_run(f"Thời gian xử lý: {stats['elapsed_s']:.2f}s\n")
        
    changes = data.get("changes", {})
    sua_doi = changes.get("sua_doi", [])
    them_moi = changes.get("them_moi", [])
    xoa_bo = changes.get("xoa_bo", [])
    
    doc.add_heading(f'I. ĐIỀU KHOẢN SỬA ĐỔI ({len(sua_doi)})', level=1)
    if not sua_doi:
        doc.add_paragraph("Không có")
    else:
        for i, item in enumerate(sua_doi, 1):
            chunk_id = item.get("vb2_chunk_id") or item.get("vb1_chunk_id") or ""
            doc.add_heading(f"{i}. {format_id(chunk_id)}", level=2)
            if item.get("summary"):
                doc.add_paragraph(f"Tóm tắt: {item['summary']}")
            
            for c in item.get("changes", []):
                if isinstance(c, dict):
                    old_p = doc.add_paragraph("Cũ: ", style="List Bullet")
                    old_p.add_run(c.get("old_content", "")).font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                    new_p = doc.add_paragraph("Mới: ", style="List Continue")
                    new_p.add_run(c.get("new_content", "")).font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                else:
                    doc.add_paragraph(str(c), style="List Bullet")
                    
    doc.add_heading(f'II. ĐIỀU KHOẢN THÊM MỚI ({len(them_moi)})', level=1)
    if not them_moi:
        doc.add_paragraph("Không có")
    else:
        for i, item in enumerate(them_moi, 1):
            chunk_id = item.get("vb2_chunk_id") or ""
            doc.add_heading(f"{i}. {format_id(chunk_id)}", level=2)
            excerpt = item.get("vb2_excerpt") or item.get("vb2", {}).get("noi_dung") or ""
            doc.add_paragraph(excerpt)

    doc.add_heading(f'III. ĐIỀU KHOẢN XÓA BỎ ({len(xoa_bo)})', level=1)
    if not xoa_bo:
        doc.add_paragraph("Không có")
    else:
        for i, item in enumerate(xoa_bo, 1):
            chunk_id = item.get("vb1_chunk_id") or ""
            doc.add_heading(f"{i}. {format_id(chunk_id)}", level=2)
            excerpt = item.get("vb1_excerpt") or item.get("vb1", {}).get("noi_dung") or ""
            doc.add_paragraph(excerpt)

    doc.add_heading('IV. TỔNG HỢP LLM', level=1)
    report_text = data.get("report_text", "")
    if report_text:
        doc.add_paragraph(report_text)
    else:
        doc.add_paragraph("Không có báo cáo.")

    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    return Path(path)

def build_html_report_content(data: Dict[str, Any]) -> str:
    stats = data.get("stats", {})
    changes = data.get("changes", {})
    sua_doi = changes.get("sua_doi", [])
    them_moi = changes.get("them_moi", [])
    xoa_bo = changes.get("xoa_bo", [])
    report_text = data.get("report_text", "")
    
    html = f"""
    <html>
    <head>
    <meta charset="utf-8">
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 40px; color: #333; line-height: 1.6; }}
        h1 {{ text-align: center; color: #1a365d; }}
        h2 {{ color: #2b6cb0; border-bottom: 2px solid #e2e8f0; padding-bottom: 5px; margin-top: 30px; }}
        table {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; table-layout: fixed; }}
        th, td {{ border: 1px solid #cbd5e0; padding: 10px; text-align: left; vertical-align: top; word-wrap: break-word; }}
        th {{ background-color: #f7fafc; }}
        .text-center {{ text-align: center; }}
        .summary {{ background: #ebf8ff; padding: 10px; border-left: 4px solid #3182ce; margin-bottom: 10px; }}
        .old-text {{ color: #c53030; background: #fff5f5; padding: 8px; margin-bottom: 5px; }}
        .new-text {{ color: #2f855a; background: #f0fff4; padding: 8px; }}
        pre {{ white-space: pre-wrap; font-family: inherit; margin: 0; }}
        .diff-block {{ margin-bottom: 10px; border: 1px solid #e2e8f0; }}
    </style>
    </head>
    <body>
        <h1>BÁO CÁO SO SÁNH VĂN BẢN PHÁP LÝ</h1>
        <p class="text-center">Ngày tạo: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')}</p>
        
        <h2>THỐNG KÊ TỔNG QUAN</h2>
        <ul>
            <li>Số lượng chunk VB1: {stats.get('so_luong_chunk_vb1', 0)}</li>
            <li>Số lượng chunk VB2: {stats.get('so_luong_chunk_vb2', 0)}</li>
            <li>Giống hệt: {stats.get('giong_nhau_hoan_toan', 0)}</li>
            <li>Sửa đổi: {stats.get('sua_doi', 0)}</li>
            <li>Thêm mới: {stats.get('them_moi', 0)}</li>
            <li>Xóa bỏ: {stats.get('xoa_bo', 0)}</li>
        </ul>
        
        <h2>I. ĐIỀU KHOẢN SỬA ĐỔI ({len(sua_doi)})</h2>
    """
    
    if sua_doi:
        html += "<table><tr><th width='5%'>STT</th><th width='25%'>Điều khoản</th><th width='70%'>Nội dung thay đổi</th></tr>"
        for i, item in enumerate(sua_doi, 1):
            chunk_id = format_id(item.get("vb2_chunk_id") or item.get("vb1_chunk_id") or "")
            summary = item.get("summary", "")
            
            changes_html = ""
            for c in item.get("changes", []):
                if isinstance(c, dict):
                    changes_html += f"<div class='diff-block'>"
                    changes_html += f"<div class='old-text'><b>Cũ:</b> <pre>{c.get('old_content', '')}</pre></div>"
                    changes_html += f"<div class='new-text'><b>Mới:</b> <pre>{c.get('new_content', '')}</pre></div>"
                    changes_html += "</div>"
                else:
                    changes_html += f"<div class='diff-block'><pre>{str(c)}</pre></div>"
                    
            html += f"""
            <tr>
                <td>{i}</td>
                <td>{chunk_id}</td>
                <td>
                    <div class="summary"><b>Tóm tắt:</b> {summary}</div>
                    {changes_html}
                </td>
            </tr>
            """
        html += "</table>"
    else:
        html += "<p>Không có</p>"
        
    html += f"<h2>II. ĐIỀU KHOẢN THÊM MỚI ({len(them_moi)})</h2>"
    if them_moi:
        html += "<table><tr><th width='10%'>STT</th><th width='25%'>Điều khoản</th><th width='65%'>Nội dung</th></tr>"
        for i, item in enumerate(them_moi, 1):
            chunk_id = format_id(item.get("vb2_chunk_id") or "")
            excerpt = item.get("vb2_excerpt") or item.get("vb2", {}).get("noi_dung") or ""
            html += f"<tr><td>{i}</td><td>{chunk_id}</td><td><pre>{excerpt}</pre></td></tr>"
        html += "</table>"
    else:
        html += "<p>Không có</p>"
        
    html += f"<h2>III. ĐIỀU KHOẢN XÓA BỎ ({len(xoa_bo)})</h2>"
    if xoa_bo:
        html += "<table><tr><th width='10%'>STT</th><th width='25%'>Điều khoản</th><th width='65%'>Nội dung</th></tr>"
        for i, item in enumerate(xoa_bo, 1):
            chunk_id = format_id(item.get("vb1_chunk_id") or "")
            excerpt = item.get("vb1_excerpt") or item.get("vb1", {}).get("noi_dung") or ""
            html += f"<tr><td>{i}</td><td>{chunk_id}</td><td><pre>{excerpt}</pre></td></tr>"
        html += "</table>"
    else:
        html += "<p>Không có</p>"
        
    html += "<h2>IV. TỔNG HỢP LLM</h2>"
    if report_text:
        html += f"<pre style='background: #f7fafc; padding: 15px; border: 1px solid #e2e8f0;'>{report_text}</pre>"
    else:
        html += "<p>Không có báo cáo.</p>"
        
    html += "</body></html>"
    return html

def build_pdf_report(data: Dict[str, Any]) -> Path:
    html_content = build_html_report_content(data)
    fd, path = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)
    
    try:
        from weasyprint import HTML
        HTML(string=html_content).write_pdf(path)
        return Path(path)
    except Exception as e:
        import logging
        logging.getLogger(__name__).warning(f"WeasyPrint failed: {e}. Falling back to HTML.")
        html_path = path + ".html"
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        return Path(html_path)
